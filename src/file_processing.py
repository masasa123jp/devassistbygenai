# -*- coding: utf-8 -*-
import streamlit as st
# file_processing.py
import os
import tempfile
import zipfile
import pandas as pd
import PyPDF2
import docx
import concurrent.futures

from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredWordDocumentLoader, UnstructuredMarkdownLoader

from src.error_handler import handle_api_errors
from src.config import TARGET_EXTENSIONS, FILE_TYPE

#ファイル別の処理定義
FILE_PROCESSORS = {
    'text/plain': lambda file: file.read().decode('utf-8') + '\n',
    'application/pdf': lambda file: '\n'.join(page.extract_text() for page in PyPDF2.PdfReader(file).pages) + '\n',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': lambda file: '\n'.join([para.text for para in docx.Document(file).paragraphs]) + '\n',
    'application/vnd.ms-excel': lambda file: read_excel_file(file),  # xls
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': lambda file: read_excel_file(file),  # xlsx
    'application/vnd.ms-excel': lambda file: read_excel_file(file),  # xlsm
    'text/x-python': lambda file: file.read().decode('utf-8') + '\n',
    'application/x-zip-compressed': lambda file: read_zip_file(file),
    'application/octet-stream': lambda file: file.read().decode('utf-8') + '\n',  # mdファイルの処理
}

# Excelファイルの読み込み処理
@handle_api_errors
def read_excel_file(file):
    content = ''
    if file.type in [FILE_TYPE['xls'], FILE_TYPE['xlsx'], FILE_TYPE['xlsm']]:
        df = pd.read_excel(file, sheet_name=None)  # すべてのシートを読み込む
        for sheet_name, sheet_data in df.items():
            content += f'シート名: {sheet_name}\n'
            content += sheet_data.to_string(index=False) + '\n\n'  # DataFrameを文字列に変換
    return content

# ZIPファイル読込処理
@handle_api_errors
def read_zip_file(uploaded_file):
    content = []
    with zipfile.ZipFile(uploaded_file, 'r') as zip_ref:
        for file in zip_ref.namelist():
            if file.lower().endswith(TARGET_EXTENSIONS):  # 対象ファイル拡張子の確認
                with zip_ref.open(file) as f:
                    if file.endswith('.txt'):
                        content.append(f.read().decode('utf-8') + '\n')
                    elif file.endswith('.pdf'):
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = '\n'.join(page.extract_text() for page in pdf_reader.pages)
                        content.append(text + '\n')
                    elif file.endswith('.docx'):
                        doc = docx.Document(f)
                        text = '\n'.join([para.text for para in doc.paragraphs])
                        content.append(text + '\n')
                    elif file.endswith('.xls') or file.endswith('.xlsx') or file.endswith('.xlsm'):
                        excel_content = read_excel_file(f)
                        content.append(excel_content)
                    elif file.endswith('.py'):
                        content.append(f.read().decode('utf-8') + '\n')
    return content

# フォルダからファイルを読み込む処理
def read_folder_files(folder_path):
    uploaded_files = []  # UploadedFileのリストを初期化

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            path = os.path.join(root, file)

            if file.lower().endswith(TARGET_EXTENSIONS):  # 対象ファイル拡張子の確認
                with open(path, 'rb') as f:
                    uploaded_files.append(f)  # UploadedFileオブジェクトを追加する処理に変更

    return uploaded_files  # list[UploadedFile]を返す

# アップロードファイルの処理
@handle_api_errors
def read_uploaded_files(uploaded_files):
    file_content = {}
    if uploaded_files:
        try:
            if len(uploaded_files) > 4:
                st.error('アップロードできるファイルの数は4ファイル以内です。')
                return file_content  # 処理を中断する
            total_size = sum(file.size for file in uploaded_files)
            if total_size > 200 * 1024 * 1024:
                st.error('アップロードファイルの総サイズは200MBを超えることはできません。')
            else:
                for uploaded_file in uploaded_files:
                    processor = FILE_PROCESSORS.get(uploaded_file.type)
                    if processor:
                        file_content[uploaded_file.name] = processor(uploaded_file)
        except Exception as e:
            st.error(f'ファイルの読み込み中にエラーが発生しました: {str(e)}')
    return file_content

def read_file(file):
    """ファイルを読み込むための補助関数"""
    try:
        with open(file, 'r') as f:
            content = f.read()
        return (file, content)  # ファイル名と内容をタプルで返す
    except Exception as e:
        return (file, str(e))  # エラーが発生した場合も返す

def read_uploaded_files_concurrent(uploaded_files):
    """アップロードされたファイルを並列に読み込む関数"""
    results = {}
    with concurrent.futures.ThreadPoolExecutor() as executor:
        # ファイルを非同期に読み込む
        future_to_file = {executor.submit(read_file, file): file for file in uploaded_files}

        for future in concurrent.futures.as_completed(future_to_file):
            file = future_to_file[future]
            try:
                file_name, content = future.result()
                results[file_name] = content
            except Exception as e:
                results[file] = str(e)

    return results

# アップロードファイルの処理を別関数に切り出す
def process_uploaded_files(uploaded_files, text_splitter):
    data = []
    for uploaded_file in uploaded_files:
        if uploaded_file.type in [FILE_TYPE['xls'], FILE_TYPE['xlsx'], FILE_TYPE['xlsm']]:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                data.extend(read_excel_file(tmp_file.name))  # Read the Excel file
        elif uploaded_file.type == FILE_TYPE['pdf']:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                loader = PyPDFLoader(tmp_file.name)
                data.extend(loader.load_and_split(text_splitter))
        elif uploaded_file.type == FILE_TYPE['txt']:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                loader = TextLoader(tmp_file.name)
                data.extend(loader.load_and_split(text_splitter))
        elif uploaded_file.type == FILE_TYPE['md']:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.md') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                loader = UnstructuredMarkdownLoader(tmp_file.name)
                data.extend(loader.load_and_split(text_splitter))
        elif uploaded_file.type == FILE_TYPE['docx']:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                loader = UnstructuredWordDocumentLoader(tmp_file.name)
                data.extend(loader.load_and_split(text_splitter))
        elif uploaded_file.type == FILE_TYPE['zip']:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.zip') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                with zipfile.ZipFile(tmp_file.name, 'r') as zip_ref:
                    for file_info in zip_ref.infolist():
                        with zip_ref.open(file_info) as file:
                            file_content = file.read()
                            if file_info.filename.endswith('.txt'):
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as txt_tmp_file:
                                    txt_tmp_file.write(file_content)
                                    loader = TextLoader(txt_tmp_file.name)
                                    data.extend(loader.load_and_split(text_splitter))
                            elif file_info.filename.endswith('.md'):
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.md') as md_tmp_file:
                                    md_tmp_file.write(file_content)
                                    loader = UnstructuredMarkdownLoader(md_tmp_file.name)
                                    data.extend(loader.load_and_split(text_splitter))
                            elif file_info.filename.endswith('.docx'):
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as docx_tmp_file:
                                    docx_tmp_file.write(file_content)
                                    loader = UnstructuredWordDocumentLoader(docx_tmp_file.name)
                                    data.extend(loader.load_and_split(text_splitter))
                            elif file_info.filename.endswith('.pdf'):
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_tmp_file:
                                    pdf_tmp_file.write(file_content)
                                    loader = PyPDFLoader(pdf_tmp_file.name)
                                    data.extend(loader.load_and_split(text_splitter))
    return data

def read_py_files_to_string(directory):
    py_files_content = ''

    # Walk through all directories and files
    for root, dirs, files in os.walk(directory):
        for file in files:
            # Check if the file is a Python (.py) file
            if file.endswith('.py'):
                file_path = os.path.join(root, file)
                try:
                    # Open and read the content of the Python file
                    with open(file_path, 'r', encoding='utf-8') as f:
                        py_files_content += f'\n# {file_path}\n'  # Add file path as a comment
                        py_files_content += f.read() + '\n'
                except Exception as e:
                    print(f'Error reading {file_path}: {e}')

    return py_files_content

# 指定されたディレクトリ配下（サブフォルダ含む）のファイルを検索し、指定された拡張子に一致するファイルの内容をUTF-8形式で読み込み返却する関数
def read_files_from_directory_as_str(directory, extensions):
    """
    指定されたディレクトリ配下（サブフォルダ含む）のファイルを検索し、
    指定された拡張子に一致するファイルの内容をUTF-8形式で読み込み、
    ファイル名とファイル内容を結合して1つの文字列として返却する。

    Args:
        directory (str): 検索を開始するトップレベルのディレクトリパス
        extensions (list): 検索するファイルの拡張子のリスト (例: ['.txt', '.csv'])

    Returns:
        str: ファイル名とファイル内容が結合された1つの文字列
    """
    combined_contents = ''

    # os.walkを使ってフォルダ内のすべてのファイルを再帰的に探索
    for root, _, files in os.walk(directory):
        for file in files:
            # 拡張子が指定されたものに一致するか確認
            if any(file.lower().endswith(ext.lower()) for ext in extensions):
                file_path = os.path.join(root, file)
                try:
                    # UTF-8形式でファイルを読み込み
                    with open(file_path, 'r', encoding='utf-8') as f:
                        file_content = f.read()
                        # ファイル名と内容を結合して追加
                        combined_contents += f'File: {file_path}  \n  '
                        combined_contents += f'Content:\n{file_content}  \n'
                        combined_contents += '-' * 40 + '\n'  # 区切りを追加
                except Exception as e:
                    print(f'Error reading file {file_path}: {e}')

    return combined_contents